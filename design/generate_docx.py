"""
Coco 导购机器人 — 硬件采购清单 Word 文档 (v6.2 双轴云台 最终版)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_cell_shading(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2C5F8A')
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)
        if r % 2 == 0:
            for c in range(len(headers)):
                set_cell_shading(table.rows[r + 1].cells[c], 'EEF2F7')
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Coco 导购机器人\n硬件采购清单 v6.2')
run.bold = True
run.font.size = Pt(30)
run.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('\nv6.2 双轴云台机型 — 总高 ~1040mm，屏幕仰角5°~55°自适应\nSG90×2舵机驱动双轴云台(pan±50° + tilt 5°~55°) + 视觉人体追踪')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('\n\n文档版本: 3.1 (v6.2)\n日期: 2026-07-04 (双轴云台 Pan+Tilt + SG90×2)\n模型文件: design/coco_model.scad\n屏幕组件: design/screen_mount.scad')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ============================================================
# 一、机器人尺寸规格
# ============================================================
doc.add_heading('一、机器人尺寸规格', level=1)

doc.add_heading('1.1 整体尺寸', level=2)
add_table(doc,
    ['项目', '数值', '说明'],
    [
        ['总高度', '~1,040 mm', '从履带接地到天线顶端'],
        ['屏幕中心高', '~840 mm', '从地面到屏幕中心'],
        ['机身宽度', '280 mm', '不含机械臂'],
        ['含臂宽度', '~540 mm', '双臂展开'],
        ['总深度', '~340 mm', '履带前后方向'],
        ['圆屏直径', '200 mm', '8寸 HDMI 圆形屏'],
        ['屏幕仰角', '5°~55°', '双轴云台自适应，28°为默认'],
        ['云台转角', 'pan ±50°, tilt 5°~55°', 'SG90×2驱动，视觉追踪'],
        ['底盘尺寸', '340 × 340 mm', '宽×长，保证稳定性'],
        ['估计重量', '12~18 kg', '含电池和所有部件'],
    ],
    col_widths=[4, 3, 7.5]
)

doc.add_paragraph()
doc.add_heading('1.2 各部件尺寸', level=2)
add_table(doc,
    ['部件', '宽 (mm)', '深 (mm)', '高 (mm)', '形状说明'],
    [
        ['履带底盘', '340', '340', '55', '圆角矩形，前后防撞条'],
        ['机身主体', '280 (底) / 220 (顶)', '220', '700', '蛋形，上窄下宽，前平后圆'],
        ['8寸圆屏', 'Ø200', '—', '—', '正圆形，5°~55°仰角自适应'],
        ['屏幕外框', 'Ø236', '—', '14', '圆环，含4个M3安装支耳+铰链耳'],
        ['头部后壳+云台', '~260', '~190', '~140', '半椭球壳+双轴云台，2.5mm壁厚'],
        ['天线总成', 'Ø32', '—', '~105', '锥形底座+杆+金色球'],
        ['单侧履带', '28', '~420', '18', '长条椭圆护罩'],
        ['驱动轮 ×4', 'Ø70', '14', '—', '6减重孔，D型轴孔'],
    ],
    col_widths=[3.5, 3, 2.5, 2.5, 4.5]
)

doc.add_paragraph()
doc.add_heading('1.3 壁厚与材料', level=2)
add_table(doc,
    ['部位', '壁厚 (mm)', '推荐材料', '填充率', '说明'],
    [
        ['机身壳体 (3段)', '3.0', 'PLA+ 白色', '15%', '非承重，美观为主'],
        ['底盘框架 (左右)', '4.0', 'PETG 灰色', '30%', '承重，需耐冲击'],
        ['屏幕外框', '2.5', 'PLA+ 白色', '20%', '含安装支耳'],
        ['头部后壳', '2.5', 'PLA+ 白色', '10%', '轻量装饰件'],
        ['履带护罩 ×2', '3.0', 'PLA+ 灰色', '20%', '保护履带'],
        ['驱动轮毂 ×4', '实心', 'PETG 灰色', '50%', '承重旋转件'],
        ['装饰件 (蝴蝶结等)', '—', 'PLA+ 深蓝 / 金色', '15%', '外观装饰'],
    ],
    col_widths=[3.5, 2, 3, 1.5, 5]
)

doc.add_paragraph()
doc.add_heading('1.4 重要: 3D打印材料说明', level=2)

p = doc.add_paragraph()
run = p.add_run('3D打印不使用 PVC。')
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph('PVC（聚氯乙烯）是做水管、电线皮的软塑料，3D打印机无法使用。常用3D打印材料如下：')
doc.add_paragraph()
add_table(doc,
    ['材料', '特点', '约价/卷', '本机用途'],
    [
        ['PLA+', '硬度好、易打印、无毒无味、表面光滑', '~¥50/kg', '机身外壳、屏幕框、装饰件'],
        ['PETG', '比PLA韧性高、耐热80°C、不易碎', '~¥60/kg', '底盘框架、轮毂'],
        ['光敏树脂 (SLA)', '表面极光滑、精度±0.1mm', '~¥80/瓶', '蝴蝶结、天线球等精细件'],
        ['ABS', '工业强度、需封闭打印机、有毒气体', '不推荐', '家用打印机不建议'],
    ],
    col_widths=[3, 6, 2, 3.5]
)

doc.add_paragraph()
doc.add_paragraph('如找淘宝代打，跟店家说: "PLA+ 白色，填充15%，壁厚3mm" 即可。底盘件加一句 "PETG 灰色，填充30%"。')

doc.add_page_break()

# ============================================================
# 二、3D 打印零件清单
# ============================================================
doc.add_heading('二、3D 打印零件清单', level=1)
doc.add_paragraph('以下为完整打印件列表。机身因700mm高度超过打印机台面，分3段拼接。')

add_table(doc,
    ['序号', '零件名', '数量', '材料', '尺寸 (mm)', '备注'],
    [
        ['1', '机身前壳 — 下段', '1', 'PLA+ 白', '280×220×230', '含屏幕倾斜凹槽+麦克风槽'],
        ['2', '机身前壳 — 中段', '1', 'PLA+ 白', '270×215×230', '上下均有定位槽'],
        ['3', '机身前壳 — 上段', '1', 'PLA+ 白', '250×210×240', '顶部收窄，连接云台底座'],
        ['4', '机身背壳 — 下左/右', '2', 'PLA+ 白', '140×220×230', '左右拼接'],
        ['5', '机身背壳 — 中左/右', '2', 'PLA+ 白', '135×215×230', '左右拼接'],
        ['6', '机身背壳 — 上左/右', '2', 'PLA+ 白', '125×210×240', '左右拼接'],
        ['7', '头部后壳+云台底座', '1', 'PLA+ 白', '260×190×140', '半球壳+双轴云台底座'],
        ['8', '屏幕外框(含铰链耳)', '1', 'PLA+ 白', 'Ø236×14', '4×M3安装孔+铰链耳，屏幕卡槽'],
        ['9', '麦克风支架', '1', 'PLA+ 白', '76×20×6', '屏幕下方'],
        ['10', '底盘框架 — 左', '1', 'PETG 灰', '340×170×55', '左侧履带+电机安装座'],
        ['11', '底盘框架 — 右', '1', 'PETG 灰', '340×170×55', '右侧履带+电机安装座'],
        ['12', '履带护罩 ×2', '2', 'PLA+ 灰', '28×420×18', '左右各一'],
        ['13', '驱动轮毂 ×4', '4', 'PETG 灰', 'Ø70×14', 'D型轴孔，6减重孔'],
        ['14', '蝴蝶结', '1', 'PLA+ 深蓝', '100×60×20', '双翼+飘带'],
        ['15', '天线底座', '1', 'PLA+ 深蓝', 'Ø32×20', '锥形'],
        ['16', '天线杆', '1', 'PLA+ 深蓝', 'Ø10×70', '细杆，实心打印'],
        ['17', '天线球', '1', 'PLA+ 金色', 'Ø34', '可喷金色漆替代'],
        ['18', '旋转底盘(pan)', '1', 'PETG 灰', 'Ø230×5', '含轴承环+舵机座'],
        ['19', '俯仰铰链支架×2', '2', 'PETG 灰', '~20×24×25', '含Ø4铰链轴孔'],
        ['20', '俯仰舵机座', '1', 'PETG 灰', '24×16×20', 'tilt SG90安装'],
        ['21', '俯仰连杆', '1', 'PLA+ 灰', '~3×15×30', '连接舵机到屏幕框'],
        ['22', '机械臂上臂 ×2', '2', 'PLA+ 灰', '含肘关节', '舵机接口'],
        ['23', '机械臂前臂 ×2', '2', 'PLA+ 灰', '含腕关节', '—'],
        ['24', '手爪 ×2', '2', 'PLA+ 红色', '三指', '—'],
        ['25', '分段连接法兰 ×4', '4', 'PLA+ 白', '内径匹配', '藏在装饰环下'],
    ],
    col_widths=[1, 3.5, 1, 2, 3, 4]
)

doc.add_paragraph()
doc.add_heading('2.1 双轴云台头部组件 STL 导出说明', level=2)
doc.add_paragraph('双轴云台(pan+tilt)已单独建模为 design/screen_mount.scad，导出步骤:')
doc.add_paragraph('打开 screen_mount.scad → 按 F6 渲染 → File → Export → Export as STL', style='List Bullet')
doc.add_paragraph('交给打印店时说: "PLA+白色(屏幕框) + PETG灰色(铰链/舵机座)，不需要支撑"', style='List Bullet')
doc.add_paragraph('屏幕外框包含4个M3安装支耳+2个铰链耳，安装时用M3螺丝锁到云台铰链支架上', style='List Bullet')
doc.add_paragraph('双轴云台: pan舵机驱动旋转底盘(±50°), tilt舵机通过连杆推动屏幕俯仰(5°~55°)', style='List Bullet')

doc.add_page_break()

# ============================================================
# 三、第一批采购 — 桌面调试
# ============================================================
doc.add_heading('三、第一批采购 — 桌面调试 (~¥1,489)', level=1)
doc.add_paragraph('目标: Orange Pi 5 + 屏幕 + 麦克风 + 喇叭, 裸板跑通完整语音对话。不需外壳和底盘。')

doc.add_heading('3.1 核心计算', level=2)
add_table(doc,
    ['序号', '名称', '规格', '数量', '单价(¥)', '淘宝搜索'],
    [
        ['1', 'Orange Pi 5', 'Orange Pi 5 4GB (RK3588)', '1', '~850', 'Orange Pi 5 4GB RK3588'],
        ['2', 'TF卡', 'SanDisk 64GB A2 V30', '1', '~45', '闪迪 64G TF卡 A2'],
        ['3', '散热套件', 'OPi 5 专用风扇+散热片', '1', '~30', 'Orange Pi 5 散热套件'],
        ['4', '电源', '5V 4A Type-C', '1', '~40', 'Orange Pi 5 电源 5V 4A'],
    ],
    col_widths=[1, 2, 3.5, 1, 1.5, 5.5]
)

doc.add_paragraph()
doc.add_heading('3.2 显示', level=2)
add_table(doc,
    ['序号', '名称', '规格', '数量', '单价(¥)', '淘宝搜索'],
    [
        ['5', '圆形屏幕', '7.9寸 HDMI 800×800 圆屏', '1', '~350',
         '7.9寸 HDMI 圆形显示屏 800*800'],
        ['备选', '方屏+遮罩', '8寸 HDMI 方屏 + 3D打印圆形遮罩面板', '1', '~150',
         '8寸 HDMI 液晶屏 800*600'],
    ],
    col_widths=[1, 2, 3.5, 1, 1.5, 5.5]
)

doc.add_paragraph()
doc.add_heading('3.3 语音', level=2)
add_table(doc,
    ['序号', '名称', '规格', '数量', '单价(¥)', '淘宝搜索'],
    [
        ['6', '麦克风阵列', 'ReSpeaker Mic Array v2.0 (USB)', '1', '~159',
         'ReSpeaker Mic Array v2.0 USB (板载3.5mm音频输出，无需额外声卡)'],
        ['7', '全频喇叭', '3W 4Ω Ø50mm', '1', '~15',
         '全频喇叭 3W 4欧 50mm'],
    ],
    col_widths=[1, 2, 3.5, 1, 1.5, 5.5]
)

doc.add_paragraph()
doc.add_heading('3.4 第一批小计', level=2)
add_table(doc,
    ['类别', '费用'],
    [['核心计算', '~¥965'], ['显示', '~¥350'], ['语音', '~¥174'], ['第一批合计', '~¥1,489']],
    col_widths=[6, 4]
)

doc.add_page_break()

# ============================================================
# 四、运动底盘 + 电源 + 连接件 + 双舵机 — 运动底盘
# ============================================================
doc.add_heading('四、第二批采购 — 运动底盘 (~¥571)', level=1)
doc.add_paragraph('目标: 底盘能前进/后退/转向, 编码器里程计可用。')

doc.add_heading('4.1 电机与驱动', level=2)
add_table(doc,
    ['序号', '名称', '规格', '数量', '单价(¥)', '淘宝搜索'],
    [
        ['9', '编码电机', 'JGA25-370 12V 100RPM (6线编码器)', '4', '~35',
         'JGA25-370 编码电机 12V 100RPM'],
        ['10', '电机驱动', 'TB6612FNG 双路驱动模块', '2', '~25',
         'TB6612FNG 电机驱动模块'],
        ['11', '橡胶履带', '内周长~800mm × 宽25mm', '2', '~40',
         '机器人橡胶履带 25mm宽'],
        ['12', '驱动轮', '金属/尼龙 Ø64mm (配履带齿距)', '4', '~15',
         '履带驱动轮 64mm D型轴'],
        ['13', '从动/张紧轮', '轴承轮 Ø64mm', '4', '~12',
         '履带张紧轮 64mm'],
        ['14', '联轴器', '6mm→8mm 弹性联轴器', '4', '~5',
         '弹性联轴器 6转8'],
        ['14+', '舵机 SG90 ×2', '微型舵机 9g 180° (双轴云台)', '2', '~8',
         'SG90 舵机 9g 180度'],
    ],
    col_widths=[1, 2, 4.2, 1, 1.5, 5]
)

doc.add_paragraph()
doc.add_heading('4.1+ 视觉追踪双轴云台', level=2)
doc.add_paragraph('SG90双舵机驱动双轴云台(pan+tilt)，配合YOLO人体检测实现自动追踪。')
doc.add_paragraph('水平追踪: 人在画面正中 → 头不动; 人偏左右 → pan舵机转头; 超出范围 → 底盘辅助转。')
doc.add_paragraph('垂直自适应: 检测人体纵向位置 → tilt舵机调仰角(TILT_MIN~TILT_MAX)。')
doc.add_paragraph('无人时: pan缓缓扫瞄搜索，tilt降到最低(5°)做出"低头失落"状态。')
doc.add_paragraph('舵机通过 Orange Pi GPIO 软件 PWM 控制，两个舵机独立线程，不需额外驱动板。')

doc.add_paragraph()
doc.add_heading('4.2 电源', level=2)
add_table(doc,
    ['序号', '名称', '规格', '数量', '单价(¥)', '淘宝搜索'],
    [
        ['15', '锂电池', '12V 6800mAh 带保护板+充电器', '1套', '~150',
         '12V 6800mAh 锂电池 保护板 充电器'],
        ['16', '降压模块', 'LM2596 DC-DC 12V→5V 5A', '1', '~15',
         'LM2596 降压模块 12V转5V'],
        ['17', '电源开关', '带灯自锁按钮 16mm', '1', '~8',
         '自锁按钮开关 16mm 带灯'],
    ],
    col_widths=[1, 2, 4.2, 1, 1.5, 5]
)

doc.add_paragraph()
doc.add_heading('4.3 连接件与工具', level=2)
add_table(doc,
    ['序号', '名称', '规格', '数量', '单价(¥)', '淘宝搜索'],
    [
        ['18', '杜邦线套装', '公母各30根 + 面包线', '1', '~15', '杜邦线 公对母 套装'],
        ['19', '铜柱螺丝', 'M2.5 六角铜柱套装 + M3螺丝套装', '1', '~15', 'M2.5/M3 铜柱 螺丝套装'],
        ['20', '热缩管', '多规格套装', '1', '~10', '热缩管 套装'],
        ['21', '扎带', '3×100mm 自锁', '1包', '~5', '扎带 3*100'],
    ],
    col_widths=[1, 2, 4.2, 1, 1.5, 5]
)

doc.add_paragraph()
doc.add_heading('4.4 第二批小计', level=2)
add_table(doc,
    ['类别', '费用'],
    [['电机与驱动', '~¥374'], ['电源', '~¥173'], ['连接件/工具', '~¥45'], ['第二批合计', '~¥592']],
    col_widths=[6, 4]
)

doc.add_page_break()

# ============================================================
# 五、第三批 — 外壳
# ============================================================
doc.add_heading('五、第三批采购 — 外壳制作', level=1)

doc.add_heading('5.1 自备3D打印机 — 买耗材', level=2)
add_table(doc,
    ['序号', '材料', '规格', '数量', '单价(¥)', '用途'],
    [
        ['22', 'PLA+ 白色', '1.75mm 1kg', '3卷', '~150', '机身(6件)+后壳(6件)+屏幕框+装饰'],
        ['23', 'PLA+ 深蓝', '1.75mm 500g', '1卷', '~30', '蝴蝶结+天线底座+装饰环'],
        ['24', 'PETG 灰色', '1.75mm 1kg', '1卷', '~60', '底盘框架+轮毂+护罩'],
    ],
    col_widths=[1, 2, 3.5, 1, 1.5, 6]
)
doc.add_paragraph('耗材合计: ~¥240')
doc.add_paragraph('注意: 机身700mm高, 分3段打印(每段≤240mm), 段间用装饰环遮盖接缝 + 内部法兰螺丝连接')

doc.add_paragraph()
doc.add_heading('5.2 无3D打印机 — 淘宝代打', level=2)
doc.add_paragraph('搜 "3D打印代打 PLA" 或 "SLA 3D打印服务", 发STL文件报价。')
doc.add_paragraph('一套全部零件代打费用约 ¥300~500 (PLA+) 到 ¥600~800 (光敏树脂高精)', style='List Bullet')
doc.add_paragraph('建议: 机身+后壳用PLA+(便宜), 蝴蝶结+天线球用光敏树脂(细节好)', style='List Bullet')

doc.add_page_break()

# ============================================================
# 六、费用总汇
# ============================================================
doc.add_heading('六、费用总汇', level=1)

add_table(doc,
    ['批次', '内容', '费用 (自打印)', '费用 (代打)'],
    [
        ['第一批', '核心计算 + 显示 + 语音', '~¥1,489', '~¥1,489'],
        ['第二批', '运动底盘 + 电源 + 连接件 + 舵机', '~¥592', '~¥592'],
        ['第三批', '外壳 (耗材/代打)', '~¥240', '~¥400~600'],
        ['', '', '', ''],
        ['合计', '全部', '~¥2,321', '~¥2,481~2,681'],
    ],
    col_widths=[2.5, 6, 3.5, 3.5]
)

doc.add_paragraph()

doc.add_page_break()

# ============================================================
# 七、购买注意事项
# ============================================================
doc.add_heading('七、购买注意事项 (避坑指南)', level=1)

doc.add_heading('7.1 圆形屏幕 — 最容易踩坑', level=2)
doc.add_paragraph('搜 "7.9寸 HDMI 圆形显示屏" 或 "Waveshare 圆形屏"')
doc.add_paragraph('确认是 HDMI 接口 (不是 MIPI)，MIPI 需要转接板才能给 Orange Pi 5 用', style='List Number')
doc.add_paragraph('分辨率 ≥ 800×800', style='List Number')
doc.add_paragraph('到货先插 Orange Pi 5 测试，Orange Pi 5 自动识别 HDMI 不需要装驱动', style='List Number')
doc.add_paragraph('7.9寸 ≈ 200mm 直径，刚好匹配模型', style='List Number')

doc.add_heading('7.2 编码电机 — 第二容易踩坑', level=2)
doc.add_paragraph('搜 "JGA25-370 编码电机 12V 100RPM"')
doc.add_paragraph('必须是 6根线 (红黑电源 + 4根编码器)，看图确认', style='List Number')
doc.add_paragraph('100RPM 是输出轴转速 (已减速)，不要买成电机本体转速', style='List Number')
doc.add_paragraph('让店家发规格书，确认安装孔距和轴径 (通常6mm D型轴)', style='List Number')
doc.add_paragraph('4个电机买同一批次，转速一致性更好', style='List Number')

doc.add_heading('7.3 履带系统 — 第三容易踩坑', level=2)
doc.add_paragraph('搜 "机器人橡胶履带 25mm宽"')
doc.add_paragraph('内周长约 800mm (配合 290mm 轮距 + 70mm 轮径)', style='List Number')
doc.add_paragraph('强烈建议: 履带 + 驱动轮 + 从动轮 同一家店买，确保齿距匹配!', style='List Number')
doc.add_paragraph('驱动轮要有 D型轴孔 或配联轴器', style='List Number')

doc.add_heading('7.4 电池安全', level=2)
doc.add_paragraph('搜 "12V 锂电池 6800mAh 保护板"')
doc.add_paragraph('必须带保护板 (过充/过放/短路保护)', style='List Number')
doc.add_paragraph('必须配专用充电器 (DC 5.5×2.1 接口, 12.6V 输出)', style='List Number')
doc.add_paragraph('蓝色热缩管包装的最常见，别买裸电芯', style='List Number')
doc.add_paragraph('充电时不要离人，锂电池安全第一', style='List Number')

doc.add_heading('7.5 3D打印 — 材料别写错', level=2)
doc.add_paragraph('3D打印不是PVC! 跟店家说 PLA+ 或 PETG (见第一章材料说明)', style='List Number')
doc.add_paragraph('家用打印机台面通常 220×220mm，机身分段已按此设计', style='List Number')
doc.add_paragraph('填充率: 机身15%，底盘30%，装饰件10%', style='List Number')
doc.add_paragraph('STL 文件导出: OpenSCAD 打开 .scad → F6 → Export as STL', style='List Number')
doc.add_paragraph('双轴云台头部组件单独导出 screen_mount.scad，含铰链耳+舵机座+屏幕框', style='List Number')

doc.add_page_break()

# ============================================================
# 八、装配顺序
# ============================================================
doc.add_heading('八、建议装配顺序', level=1)

steps = [
    ('第1步: 桌面调试 (第一批到货后)',
     'Pi + 屏幕 + 麦克风 + 喇叭，裸板接通。\n'
     '安装依赖: pip install sounddevice openai-whisper ...\n'
     '跑 voice_pipeline.py 自测: 说"你好Coco" → 语音回复。\n'
     '屏幕显示 Coco 表情 (运行 UI 模块)。'),
    ('第2步: 底盘搭建 (第二批到货后)',
     '电机 + TB6612 + 履带 + 电池，裸底盘先跑。\n'
     '测试: 前进/后退/原地转向, 编码器读数正确。\n'
     'PID 速度环标定 (config.py 里的 PID 参数需实车调)。'),
    ('第3步: 外壳安装 (第三批到货后)',
     '打印件到货 → 打磨毛刺 → 先干装确认孔位对齐。\n'
     '双轴云台: 铰链支架→旋转平台, tilt舵机座→平台, screen_mount→铰链轴+连杆。\n'
     '屏幕: 卡入 screen_mount → M3×8mm螺丝 ×4锁到铰链耳。\n'
     '喇叭: Ø50mm全频喇叭贴到格栅背面 → M3×12mm螺丝 ×2固定到speaker_mount。\n'
     '机身分段: 法兰螺丝对接 + 装饰环遮缝。\n'
     '底盘框架: M4螺丝固定到机身底部。'),
    ('第4步: 总装联调',
     '全部电子件塞入机身，扎带理线，电池固定在底盘。\n'
     '电机线从底部出线孔 → 驱动板 → Orange Pi 5 GPIO。\n'
     '上电跑完整场景: 屏幕亮表情 → 麦克风拾音 → LLM对话 → TTS播报 → 移动。'),
]

for title, desc in steps:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

doc.add_page_break()

# ============================================================
# 九、设计图纸索引
# ============================================================
doc.add_heading('九、设计图纸与模型文件', level=1)

add_table(doc,
    ['文件', '说明', '用途'],
    [
        ['design/coco_model.scad', 'OpenSCAD 3D模型 v6.2', 'F6渲染后导出全部STL'],
        ['design/screen_mount.scad', '双轴云台头部组件(pan+tilt)', '独立导出头部STL（含铰链+舵机座）'],
        ['design/coco_concept.png', '正面概念效果图 (1200×1900)', '参考外观'],
        ['design/coco_side.png', '侧面剖视图 (1600×1200)', '参考侧面轮廓+屏幕仰角'],
        ['design/screen_mount_guide.png', '屏幕安装示意图', '参考安装方式'],
        ['design/Coco硬件采购清单.docx', '本文档', '采购+装配指南'],
        ['config.py', '全局配置文件', '电机/音频/LLM/TTS 所有参数'],
    ],
    col_widths=[5, 4.5, 4]
)

doc.add_paragraph()
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('— END —')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ===== 保存 =====
output = 'D:/claude/coco/design/Coco硬件采购清单.docx'
doc.save(output)
print(f'Word 已保存: {output}')
print(f'大小: {os.path.getsize(output)/1024:.1f} KB')
